import subprocess
import platform
import socket
import speedtest
import requests
from rich.console import Console
from rich.table import Table
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TaskProgressColumn
from rich.panel import Panel
from rich.columns import Columns
from rich.status import Status
from rich.align import Align
from rich.text import Text
from docx import Document
from datetime import datetime
import argparse
import json
import csv
import matplotlib.pyplot as plt
from docx.shared import Inches
import os
import re
from fpdf import FPDF
import time
console = Console()

# Configuração de estilo para matplotlib
try:
    import seaborn as sns
    HAS_SEABORN = True
    plt.style.use('seaborn-v0_8')
except ImportError:
    HAS_SEABORN = False
    plt.style.use('default')

DESTINOS = {
    "🌐 Google": "google.com",
    "🤖 OpenAI": "openai.com", 
    "💻 GitHub": "github.com",
    "☁️ Cloudflare DNS": "1.1.1.1",
    "🔍 Google DNS": "8.8.8.8",
}

# Configuração dos provedores
PROVEDORES = {
    "Vivo": {
        "ips_conhecidos": [
            "201.95.209.43", "201.95.215.228", "187.57.45.67"
        ],
        "ranges": ["201.95.", "187.57."],
        "hostname_pattern": r".*\.dsl\.telesp\.net\.br",
        "tipo": "IP Dinâmico"
    },
    "Netflex": {
        "ips_conhecidos": ["170.79.166.98"],
        "ranges": ["170.79.166."],
        "hostname_pattern": r".*netflex.*",
        "tipo": "IP Fixo"
    }
}

def obter_ip_publico():
    """Obtém o IP público atual usando múltiplos serviços"""
    servicos = [
        "https://api.ipify.org",
        "https://ifconfig.me/ip",
        "https://icanhazip.com",
        "https://checkip.amazonaws.com"
    ]
    
    for servico in servicos:
        try:
            response = requests.get(servico, timeout=5)
            if response.status_code == 200:
                ip = response.text.strip()
                # Validar se é um IP válido
                if re.match(r'^\d+\.\d+\.\d+\.\d+$', ip):
                    return ip
        except:
            continue
    
    return None

def obter_hostname_reverso(ip):
    """Obtém o hostname reverso do IP"""
    try:
        hostname = socket.gethostbyaddr(ip)[0]
        return hostname
    except:
        return None

def detectar_provedor():
    """Detecta qual provedor está sendo usado"""
    ip_publico = obter_ip_publico()
    
    if not ip_publico:
        return {
            "provedor": "Desconhecido",
            "ip": "Não detectado",
            "hostname": None,
            "tipo": "N/A",
            "confianca": 0
        }
    
    hostname = obter_hostname_reverso(ip_publico)
    
    # Verificar cada provedor
    for nome_provedor, config in PROVEDORES.items():
        confianca = 0
        
        # Verificar IP exato (maior confiança)
        if ip_publico in config["ips_conhecidos"]:
            confianca += 90
        
        # Verificar range de IPs
        for range_ip in config["ranges"]:
            if ip_publico.startswith(range_ip):
                confianca += 70
                break
        
        # Verificar hostname pattern
        if hostname and re.match(config["hostname_pattern"], hostname, re.IGNORECASE):
            confianca += 60
        
        # Se confiança for alta o suficiente, retornar este provedor
        if confianca >= 60:
            return {
                "provedor": nome_provedor,
                "ip": ip_publico,
                "hostname": hostname,
                "tipo": config["tipo"],
                "confianca": confianca
            }
    
    # Se não identificou nenhum provedor conhecido
    return {
        "provedor": "Desconhecido",
        "ip": ip_publico,
        "hostname": hostname,
        "tipo": "N/A",
        "confianca": 0
    }

def mostrar_info_conexao():
    """Exibe informações sobre a conexão atual"""
    with Status("[cyan]🔍 Detectando provedor de internet...", spinner="dots"):
        info_conexao = detectar_provedor()
    
    # Determinar cor baseada na confiança
    if info_conexao["confianca"] >= 90:
        cor_provedor = "green"
        icone_confianca = "🟢"
    elif info_conexao["confianca"] >= 60:
        cor_provedor = "yellow"
        icone_confianca = "🟡"
    else:
        cor_provedor = "red"
        icone_confianca = "🔴"
    
    # Criar tabela de informações da conexão
    table = Table(show_header=True, header_style="bold cyan", width=100)
    table.add_column("🌐 Informação", style="cyan", width=25)
    table.add_column("📋 Valor", style="white", width=35)
    table.add_column("📊 Status", justify="center", width=10)
    
    # Ícone do provedor
    icone_provedor = "📶" if info_conexao["provedor"] != "Desconhecido" else "❓"
    
    table.add_row(
        f"{icone_provedor} Provedor", 
        f"[{cor_provedor}]{info_conexao['provedor']}[/{cor_provedor}]", 
        icone_confianca
    )
    table.add_row("🌍 IP Público", info_conexao["ip"], "🔍")
    table.add_row("🏷️ Tipo de IP", info_conexao["tipo"], "📋")
    
    if info_conexao["hostname"]:
        table.add_row("🖥️ Hostname", info_conexao["hostname"][:45], "🔗")
    
    # Mostrar confiança apenas se for identificado
    if info_conexao["confianca"] > 0:
        table.add_row("🎯 Confiança", f"{info_conexao['confianca']}%", icone_confianca)
    
    console.print("\n")
    console.print(Panel(table, title="🔌 Informações da Conexão", border_style="blue"))
    
    return info_conexao

def mostrar_banner():
    """Exibe um banner de boas-vindas atrativo"""
    banner_text = Text("🌐 DIAGNÓSTICO DE REDE 🌐", style="bold cyan")
    subtitle_text = Text("Análise completa de conectividade e performance", style="italic blue")
    
    banner_panel = Panel(
        Align.center(f"{banner_text}\n{subtitle_text}"), 
        border_style="blue",
        padding=(1, 2)
    )
    console.print(banner_panel)
    console.print()

def criar_tabela_dns(resultados_dns):
    """Cria uma tabela organizada para os resultados DNS"""
    table = Table(show_header=True, header_style="bold magenta", width=120)
    table.add_column("🌐 Serviço", style="cyan", no_wrap=True, width=20)
    table.add_column("🔍 Host", style="white", width=15)
    table.add_column("📍 IP Resolvido", style="green", width=18)
    table.add_column("Status", justify="center", width=8)
    
    for nome, host in DESTINOS.items():
        ip = resultados_dns.get(f"DNS {nome.split(' ', 1)[-1]}", "Não resolvido")
        status = "✅" if not ip.startswith("Erro") else "❌"
        table.add_row(nome, host, ip, status)
    
    return table

def criar_tabela_speedtest(speed_result):
    """Cria uma tabela atrativa para os resultados do speedtest"""
    table = Table(show_header=True, header_style="bold yellow", width=80)
    table.add_column("📊 Métrica", style="cyan", width=18)
    table.add_column("📈 Valor", style="green", justify="right", width=12)
    table.add_column("📋 Unidade", style="blue", width=10)
    table.add_column("🎯 Status", justify="center", width=8)
    
    down, up, ping_ms = speed_result
    
    # Status baseado em thresholds típicos
    down_status = "🟢" if down > 50 else "🟡" if down > 10 else "🔴"
    up_status = "🟢" if up > 10 else "🟡" if up > 2 else "🔴"
    ping_status = "🟢" if ping_ms < 50 else "🟡" if ping_ms < 100 else "🔴"
    
    table.add_row("📥 Download", f"{down:.2f}", "Mbps", down_status)
    table.add_row("📤 Upload", f"{up:.2f}", "Mbps", up_status)
    table.add_row("⚡ Latência", f"{ping_ms:.2f}", "ms", ping_status)
    
    return table

def ping(host):
    param = "-n" if platform.system().lower() == "windows" else "-c"
    cmd = ["ping", param, "4", host]
    return subprocess.run(cmd, stdout=subprocess.PIPE).stdout.decode()

def traceroute(host):
    cmd = ["tracert" if platform.system().lower() == "windows" else "traceroute", host]
    return subprocess.run(cmd, stdout=subprocess.PIPE).stdout.decode()

def resolve_dns(host):
    try:
        return socket.gethostbyname(host)
    except socket.error as e:
        return f"Erro: {e}"

def formatar_ping_resumo(resultado_ping):
    """Extrai informações resumidas do resultado do ping"""
    linhas = resultado_ping.strip().split('\n')
    
    # Extrair informações básicas
    info = {
        'host': '',
        'ip': '',
        'pacotes_enviados': 0,
        'pacotes_recebidos': 0,
        'perda': '0%',
        'latencia_min': 0,
        'latencia_avg': 0,
        'latencia_max': 0
    }
    
    # Primeira linha com host e IP
    primeira_linha = linhas[0] if linhas else ""
    if "PING" in primeira_linha:
        partes = primeira_linha.split()
        if len(partes) >= 3:
            info['host'] = partes[1]
            # Extrair IP entre parênteses
            ip_match = re.search(r'\(([^)]+)\)', primeira_linha)
            if ip_match:
                info['ip'] = ip_match.group(1)
    
    # Estatísticas finais
    for linha in reversed(linhas):
        if "packet loss" in linha:
            # Extrair estatísticas de perda
            match = re.search(r'(\d+) packets transmitted, (\d+) received, (\d+(?:\.\d+)?%)', linha)
            if match:
                info['pacotes_enviados'] = int(match.group(1))
                info['pacotes_recebidos'] = int(match.group(2))
                info['perda'] = match.group(3)
        elif "rtt min/avg/max" in linha:
            # Extrair latências
            match = re.search(r'= ([\d\.]+)/([\d\.]+)/([\d\.]+)/([\d\.]+) ms', linha)
            if match:
                info['latencia_min'] = float(match.group(1))
                info['latencia_avg'] = float(match.group(2))
                info['latencia_max'] = float(match.group(3))
    
    return info

def formatar_traceroute_resumo(resultado_traceroute):
    """Extrai rota simplificada do traceroute"""
    linhas = resultado_traceroute.strip().split('\n')
    
    rota = []
    destino_info = ""
    
    # Primeira linha com destino
    if linhas and "traceroute to" in linhas[0]:
        destino_info = linhas[0]
    
    # Processar cada salto
    for linha in linhas[1:]:
        if re.match(r'^\s*\d+', linha):
            hop_num = re.match(r'^\s*(\d+)', linha).group(1)
            
            # Extrair IPs e hostnames
            hosts = []
            tempos = []
            
            # Pegar todos os IPs/hostnames da linha
            ip_pattern = r'(\d+\.\d+\.\d+\.\d+)'
            hostname_pattern = r'([a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
            tempo_pattern = r'(\d+(?:\.\d+)?)\s*ms'
            
            ips = re.findall(ip_pattern, linha)
            hostnames = re.findall(hostname_pattern, linha)
            tempos_found = re.findall(tempo_pattern, linha)
            
            # Pegar primeiro hostname válido ou IP
            host_principal = ""
            if hostnames:
                # Filtrar hostnames que não são IPs
                hostnames_validos = [h for h in hostnames if not re.match(r'^\d+\.\d+\.\d+\.\d+$', h)]
                if hostnames_validos:
                    host_principal = hostnames_validos[0]
            
            if not host_principal and ips:
                host_principal = ips[0]
            
            # Se não tem host, é timeout
            if not host_principal and '*' in linha:
                host_principal = "* timeout"
            
            # Latência média dos tempos encontrados
            latencia_media = 0
            if tempos_found:
                latencias = [float(t) for t in tempos_found]
                latencia_media = sum(latencias) / len(latencias)
            
            if host_principal:
                rota.append({
                    'hop': int(hop_num),
                    'host': host_principal,
                    'latencia': round(latencia_media, 1) if latencia_media > 0 else 0
                })
    
    return {
        'destino_info': destino_info,
        'rota': rota
    }

def formatar_mtr_resumo(resultado_mtr):
    """Extrai resumo do MTR focando nos pontos críticos"""
    linhas = resultado_mtr.strip().split('\n')
    
    resumo = {
        'destino': '',
        'hops_problematicos': [],
        'perda_total': 0,
        'latencia_final': 0
    }
    
    # Extrair destino
    for linha in linhas:
        if "HOST:" in linha:
            parts = linha.split()
            if len(parts) > 1:
                resumo['destino'] = parts[1]
            break
    
    # Processar linhas de dados
    for linha in linhas:
        if re.match(r'^\s*\d+\.\|--', linha):
            partes = linha.split()
            if len(partes) >= 6:
                try:
                    hop_num = int(partes[0].replace('.', '').replace('|--', ''))
                    host = partes[1]
                    loss_pct = float(partes[2].replace('%', ''))
                    
                    # Ignorar linhas com ???
                    if '???' in host:
                        continue
                    
                    # Se tem perda significativa, adicionar aos problemáticos
                    if loss_pct > 0:
                        resumo['hops_problematicos'].append({
                            'hop': hop_num,
                            'host': host,
                            'perda': loss_pct
                        })
                    
                    # Última linha válida para latência final
                    if len(partes) >= 7 and loss_pct < 100:
                        try:
                            latencia = float(partes[5])  # Coluna Avg
                            resumo['latencia_final'] = latencia
                        except:
                            pass
                            
                except:
                    continue
    
    return resumo

def test_speed(server_id=None):
    try:
        st = speedtest.Speedtest(timeout=20)
        st.get_servers()
        if server_id:
            servers = [s for s in st.servers[server_id]]
            st.get_best_server(servers)
        else:
            st.get_best_server()
        down = st.download() / 1e6
        up = st.upload() / 1e6
        ping = st.results.ping
        return round(down, 2), round(up, 2), round(ping, 2)
    except Exception as e:
        return (0, 0, 0), f"Erro no speedtest: {e}"

def gerar_pasta_resultados():
    pasta = datetime.now().strftime("resultados_%Y%m%d_%H%M%S")
    os.makedirs(pasta, exist_ok=True)
    return pasta

def exportar_json(resultados, pasta, nome_arquivo=None):
    """Exporta resultados para JSON com feedback visual melhorado"""
    if nome_arquivo is None:
        nome_arquivo = f"resultados_{pasta}.json"
    caminho = os.path.join(pasta, nome_arquivo)
    with open(caminho, "w", encoding='utf-8') as f:
        json.dump(resultados, f, indent=4, ensure_ascii=False)
    console.print(f"✅ [bold green]JSON exportado:[/] [cyan]{caminho}[/cyan]")

def exportar_csv(resultados, pasta, nome_arquivo=None):
    """Exporta resultados para CSV com feedback visual melhorado"""
    if nome_arquivo is None:
        nome_arquivo = f"resultados_{pasta}.csv"
    caminho = os.path.join(pasta, nome_arquivo)
    with open(caminho, "w", newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(["Teste", "Resultado"])
        for nome, dados in resultados.items():
            writer.writerow([nome, str(dados)])
    console.print(f"✅ [bold green]CSV exportado:[/] [cyan]{caminho}[/cyan]")

def gerar_grafico_speedtest(speed, nome_imagem):
    """Gera um gráfico moderno e atrativo para o speedtest"""
    # Usar labels sem emojis para compatibilidade com PDF
    if "pdf" in nome_imagem.lower() or "metricas" in nome_imagem.lower():
        labels = ["Download", "Upload", "Latência"]
    else:
        labels = ["📥 Download", "📤 Upload", "⚡ Latência"]
    
    valores = [speed[0], speed[1], speed[2]]
    
    # Cores modernas em gradiente
    cores = ['#00d4aa', '#ff6b6b', '#4ecdc4']
    
    plt.figure(figsize=(10, 6))
    
    # Criando o gráfico de barras com estilo moderno
    bars = plt.bar(labels, valores, color=cores, alpha=0.8, edgecolor='white', linewidth=2)
    
    # Adicionando valores no topo das barras
    for bar, valor in zip(bars, valores):
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height + height*0.01,
                f'{valor:.1f}', ha='center', va='bottom', fontweight='bold', fontsize=12)
    
    # Estilização do gráfico
    plt.title('Teste de Velocidade da Internet', fontsize=16, fontweight='bold', pad=20)
    plt.ylabel('Velocidade (Mbps) / Latência (ms)', fontsize=12)
    plt.grid(axis='y', alpha=0.3, linestyle='--')
    
    # Melhorando a aparência
    plt.gca().spines['top'].set_visible(False)
    plt.gca().spines['right'].set_visible(False)
    plt.gca().set_facecolor('#f8f9fa')
    
    plt.tight_layout()
    plt.savefig(nome_imagem, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return nome_imagem

def gerar_doc(resultados, pasta):
    doc = Document()
    doc.add_heading("Relatório de Diagnóstico de Rede", level=1)
    doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")

    # Resumo dos resultados
    doc.add_heading("Resumo", level=2)
    resumo = []
    for nome, dados in resultados.items():
        if nome == "Speedtest":
            continue
        if isinstance(dados, str) and dados.startswith("Erro"):
            resumo.append(f"{nome}: {dados}")
        else:
            resumo.append(f"{nome}: OK")
    speed = resultados.get("Speedtest", (0, 0, 0))
    resumo.append(f"Speedtest: Download {speed[0]} Mbps, Upload {speed[1]} Mbps, Latência {speed[2]} ms")
    doc.add_paragraph("\n".join(resumo), style='Normal')
    doc.add_paragraph("\n")

    # Índice (sumário simples)
    doc.add_heading("Índice", level=2)
    for nome in resultados.keys():
        if nome == "Speedtest":
            continue
        doc.add_paragraph(nome, style='List Number')
    doc.add_paragraph("Speedtest", style='List Number')
    doc.add_paragraph("\n")

    # Resultados detalhados
    for nome, dados in resultados.items():
        if nome == "Speedtest":
            continue
        doc.add_heading(nome, level=2)
        doc.add_paragraph("\n")
        if isinstance(dados, str):
            doc.add_paragraph(dados, style='Normal')
        else:
            doc.add_paragraph(str(dados), style='Normal')
        doc.add_paragraph("\n")

    doc.add_heading("Speedtest", level=2)
    doc.add_paragraph("\n")
    doc.add_paragraph(
        f"Download: {speed[0]} Mbps\nUpload: {speed[1]} Mbps\nLatência: {speed[2]} ms",
        style='Intense Quote'
    )

    # Adiciona gráfico ao Word
    nome_imagem = gerar_grafico_speedtest(resultados.get("Speedtest", (0,0,0)), os.path.join(pasta, f"grafico_speedtest_{pasta}.png"))
    doc.add_picture(nome_imagem, width=Inches(4.5))
    doc.add_paragraph("\n")

    nome_arquivo = os.path.join(pasta, f"Relatorio_Diagnostico_Rede_{pasta}.docx")
    doc.save(nome_arquivo)
    console.print(f"[bold green]Relatório salvo como:[/] {nome_arquivo}")

def mtr_analysis(host):
    # Executa mtr em modo relatório, 10 pacotes
    cmd = ["mtr", "--report", "--report-cycles", "10", host]
    try:
        resultado = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
        return resultado.stdout.decode()
    except subprocess.CalledProcessError as e:
        return f"Erro ao executar mtr: {e.stderr.decode()}"

def log(msg, logfile):
    with open(logfile, "a") as f:
        f.write(msg + "\n")

def extrair_resultados_por_destino(log_path):
    resultados = {}
    with open(log_path, "r") as f:
        log = f.read()

    # Mapear nomes com emojis para nomes simples
    nomes_simples = {
        "🌐 Google": "Google",
        "🤖 OpenAI": "OpenAI", 
        "💻 GitHub": "GitHub",
        "☁️ Cloudflare DNS": "Cloudflare DNS",
        "🔍 Google DNS": "Google DNS",
    }

    for nome_emoji, host in DESTINOS.items():
        nome_simples = nomes_simples[nome_emoji]
        
        # Extrai Ping
        ping_match = re.search(rf"^Ping {nome_simples}:([\s\S]*?)(?=^Ping|^Traceroute|^MTR|$)", log, re.MULTILINE)
        ping_result = ping_match.group(0).strip() if ping_match else "Não encontrado"

        # Extrai Traceroute
        traceroute_match = re.search(rf"^Traceroute {nome_simples}:([\s\S]*?)(?=^Ping|^Traceroute|^MTR|$)", log, re.MULTILINE)
        traceroute_result = traceroute_match.group(0).strip() if traceroute_match else "Não encontrado"

        # Extrai MTR
        mtr_match = re.search(rf"^MTR {nome_simples}:([\s\S]*?)(?=^Ping|^Traceroute|^MTR|$)", log, re.MULTILINE)
        mtr_result = mtr_match.group(0).strip() if mtr_match else "Não encontrado"

        resultados[nome_simples] = {
            "Ping": ping_result,
            "Traceroute": traceroute_result,
            "MTR": mtr_result
        }
    return resultados

def gerar_pdf_resultados_por_destino(resultados, speedtest, caminho_saida_pdf="Relatorio_Diagnostico_Rede.pdf"):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Uma página por destino
    for nome, dados in resultados.items():
        pdf.add_page()
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, f"Destino: {nome}", ln=True)
        pdf.set_font("Arial", '', 12)
        pdf.ln(5)
        pdf.multi_cell(0, 8, f"Ping:\n{dados['Ping']}\n")
        pdf.ln(2)
        pdf.multi_cell(0, 8, f"Traceroute:\n{dados['Traceroute']}\n")
        pdf.ln(2)
        pdf.multi_cell(0, 8, f"MTR:\n{dados['MTR']}\n")

    # Página final: Speedtest
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, "Speedtest", ln=True)
    pdf.set_font("Arial", '', 12)
    pdf.ln(5)
    pdf.cell(0, 10, f"Download: {speedtest[0]} Mbps", ln=True)
    pdf.cell(0, 10, f"Upload: {speedtest[1]} Mbps", ln=True)
    pdf.cell(0, 10, f"Latência: {speedtest[2]} ms", ln=True)

    # Gráfico do Speedtest
    labels = ["Download", "Upload", "Latência"]
    valores = [speedtest[0], speedtest[1], speedtest[2]]
    plt.figure(figsize=(6,4))
    plt.bar(labels, valores, color=["blue", "green", "orange"])
    plt.ylabel("Mbps / ms")
    plt.title("Speedtest")
    plt.tight_layout()
    grafico_path = "grafico_speedtest.png"
    plt.savefig(grafico_path)
    plt.close()
    pdf.image(grafico_path, w=120)

    pdf.output(caminho_saida_pdf)
    print(f"[✔] PDF salvo como {caminho_saida_pdf}")

def gerar_pdf_metricas_por_destino(resultados, speedtest, pasta, caminho_saida_pdf="Relatorio_Diagnostico_Rede_Metricas.pdf"):
    """Gera PDF com métricas e gráficos modernos"""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Cores modernas em hex
    cores_modernas = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FECA57', '#FF9FF3']

    for i, (nome, dados) in enumerate(resultados.items()):
        pdf.add_page()
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, f"Destino: {nome}", ln=True)
        pdf.set_font("Arial", '', 12)
        pdf.ln(5)
        
        for k, v in dados.items():
            pdf.cell(0, 10, f"{k}: {v}", ln=True)

        # Gráfico moderno para cada destino
        labels = list(dados.keys())
        valores = list(dados.values())
        
        plt.figure(figsize=(10, 6))
        bars = plt.bar(labels, valores, color=cores_modernas[:len(labels)], alpha=0.8, edgecolor='white', linewidth=2)
        
        # Adicionando valores nas barras
        for bar, valor in zip(bars, valores):
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., height + height*0.01,
                    f'{valor}', ha='center', va='bottom', fontweight='bold', fontsize=10)
        
        plt.title(f"Métricas de Rede - {nome}", fontsize=14, fontweight='bold', pad=20)
        plt.ylabel("Valor", fontsize=12)
        plt.xticks(rotation=30, ha='right')
        plt.grid(axis='y', alpha=0.3, linestyle='--')
        
        # Melhorando aparência
        plt.gca().spines['top'].set_visible(False)
        plt.gca().spines['right'].set_visible(False)
        plt.gca().set_facecolor('#f8f9fa')
        
        plt.tight_layout()
        grafico_path = os.path.join(pasta, f"grafico_metricas_{nome.replace(' ', '_')}.png")
        plt.savefig(grafico_path, dpi=200, bbox_inches='tight', facecolor='white')
        plt.close()
        
        pdf.image(grafico_path, w=150)
        pdf.ln(10)

    # Página do Speedtest com design melhorado
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, "Speedtest", ln=True)
    pdf.set_font("Arial", '', 12)
    pdf.ln(5)
    pdf.cell(0, 10, f"Download: {speedtest[0]} Mbps", ln=True)
    pdf.cell(0, 10, f"Upload: {speedtest[1]} Mbps", ln=True)
    pdf.cell(0, 10, f"Latencia: {speedtest[2]} ms", ln=True)

    # Gráfico moderno do speedtest
    labels = ["Download", "Upload", "Latencia"]
    valores = [speedtest[0], speedtest[1], speedtest[2]]
    cores_speed = ['#00d4aa', '#ff6b6b', '#4ecdc4']
    
    plt.figure(figsize=(10, 6))
    bars = plt.bar(labels, valores, color=cores_speed, alpha=0.8, edgecolor='white', linewidth=2)
    
    for bar, valor in zip(bars, valores):
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height + height*0.01,
                f'{valor:.1f}', ha='center', va='bottom', fontweight='bold', fontsize=12)
    
    plt.title('Teste de Velocidade da Internet', fontsize=16, fontweight='bold', pad=20)
    plt.ylabel('Velocidade (Mbps) / Latência (ms)', fontsize=12)
    plt.grid(axis='y', alpha=0.3, linestyle='--')
    
    plt.gca().spines['top'].set_visible(False)
    plt.gca().spines['right'].set_visible(False)
    plt.gca().set_facecolor('#f8f9fa')
    
    plt.tight_layout()
    grafico_path = os.path.join(pasta, "grafico_speedtest_moderno.png")
    plt.savefig(grafico_path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    
    pdf.image(grafico_path, w=150)

    pdf.output(os.path.join(pasta, caminho_saida_pdf))
    console.print(f"[bold green]✅ PDF de métricas salvo como:[/] {os.path.join(pasta, caminho_saida_pdf)}")

def extrair_metricas_do_log(log_path):
    resultados = {}
    with open(log_path, "r") as f:
        log = f.read()

    # Mapear nomes com emojis para nomes simples
    nomes_simples = {
        "🌐 Google": "Google",
        "🤖 OpenAI": "OpenAI", 
        "💻 GitHub": "GitHub",
        "☁️ Cloudflare DNS": "Cloudflare DNS",
        "🔍 Google DNS": "Google DNS",
    }

    for nome_emoji, host in DESTINOS.items():
        nome_simples = nomes_simples[nome_emoji]
        
        # Extrai o bloco do traceroute desse destino - ajustado para não parar em "Traceroute X DNS"
        traceroute_match = re.search(
            rf"^Traceroute {nome_simples}:([\s\S]*?)(?=^(?:Ping|Traceroute (?!{nome_simples})|MTR|Speedtest))", log, re.MULTILINE)
        traceroute_bloco = traceroute_match.group(1) if traceroute_match else ""

        # Ping (latência média) - melhor correspondência para evitar conflitos
        ping_match = re.search(
            rf"Ping {re.escape(nome_simples)}.*?rtt min/avg/max/mdev = ([\d\.]+)/([\d\.]+)/([\d\.]+)/([\d\.]+) ms",
            log, re.DOTALL)
        if ping_match:
            latencia = float(ping_match.group(2))
        else:
            # Se não encontrou estatísticas, pode ser 100% de perda (como PlatformIO)
            if re.search(rf"Ping {re.escape(nome_simples)}.*?100% packet loss", log, re.DOTALL):
                latencia = 0.0  # Sem resposta
            else:
                latencia = 0.0

        # Traceroute: extrair latência final - pegar todos os tempos das linhas
        # Buscar por linhas que têm tempos em ms, incluindo casos com múltiplos IPs na mesma linha
        # Filtrar linhas com apenas asteriscos
        tempos_todas_linhas = []
        linhas_traceroute = traceroute_bloco.split('\n')
        
        for linha in linhas_traceroute:
            # Ignorar linhas que são apenas asteriscos (ex: "3  * * *")
            if re.match(r'^\s*\d+\s+\*\s*\*\s*\*\s*$', linha.strip()):
                continue
            # Pegar todos os tempos em ms de cada linha
            tempos_linha = re.findall(r'(\d+\.\d+) ms', linha)
            if tempos_linha:
                tempos_todas_linhas.extend([float(t) for t in tempos_linha])
        
        # Pegar os últimos 3 tempos como latência final
        if len(tempos_todas_linhas) >= 3:
            lat_final = sum(tempos_todas_linhas[-3:]) / 3
        elif tempos_todas_linhas:
            lat_final = sum(tempos_todas_linhas) / len(tempos_todas_linhas)
        else:
            lat_final = 0.0

        # Perda intermediária traceroute: contar saltos com asteriscos
        linhas_traceroute_numeradas = re.findall(r'^\s*\d+\s+.*$', traceroute_bloco, re.MULTILINE)
        saltos_com_perda = 0
        total_saltos = len(linhas_traceroute_numeradas)
        
        for linha in linhas_traceroute_numeradas:
            # Contar quantos asteriscos há na linha (indicando timeout)
            if '*' in linha:
                saltos_com_perda += 1
        
        perda_pct = round(100 * saltos_com_perda / total_saltos, 1) if total_saltos > 0 else 0.0

        # MTR (média de perda intermediária) - por destino - ajustado para não parar em "MTR X DNS"
        mtr_match = re.search(
            rf"^MTR {nome_simples}:([\s\S]*?)(?=^(?:Ping|Traceroute|MTR (?!{nome_simples})|Speedtest))", log, re.MULTILINE)
        mtr_bloco = mtr_match.group(1) if mtr_match else ""
        
        # Extrair perdas do MTR - formato: "|-- host Loss% ..."
        # Filtrar linhas com "???" que indicam hosts não identificados
        perdas_mtr = []
        linhas_mtr = mtr_bloco.split('\n')
        for linha in linhas_mtr:
            # Ignorar linhas que contêm "???" (hosts não identificados)
            if '???' in linha:
                continue
            # Buscar percentual de perda na linha
            match_perda = re.search(r'\|\-\-\s+\S+.*?(\d+\.\d+)%', linha)
            if match_perda:
                perda_valor = float(match_perda.group(1))
                if perda_valor > 0:  # Só considerar perdas > 0
                    perdas_mtr.append(perda_valor)
        
        if perdas_mtr:
            perda_intermediaria = sum(perdas_mtr) / len(perdas_mtr)
        else:
            perda_intermediaria = 0.0

        resultados[nome_simples] = {
            "latência média ping (ms)": latencia,
            "latência final traceroute (ms)": round(lat_final, 2),
            "perda intermediária traceroute (%)": perda_pct,
            "perda intermediária MTR (%)": round(perda_intermediaria, 1)
        }
    return resultados

def main():
    parser = argparse.ArgumentParser(description="🌐 Diagnóstico Completo de Rede")
    parser.add_argument("--speedtest-server", type=int, help="ID do servidor do Speedtest")
    args = parser.parse_args()

    # Banner de boas-vindas
    mostrar_banner()
    
    # Detectar e mostrar informações da conexão
    info_conexao = mostrar_info_conexao()
    
    # Criar pasta de resultados primeiro
    pasta = gerar_pasta_resultados()
    
    resultados = {}
    # Adicionar informações da conexão aos resultados
    resultados["Conexao"] = {
        "provedor": info_conexao["provedor"],
        "ip_publico": info_conexao["ip"],
        "hostname": info_conexao["hostname"],
        "tipo_ip": info_conexao["tipo"],
        "confianca_deteccao": info_conexao["confianca"]
    }
    
    # Log será salvo dentro da pasta de resultados
    logfile = os.path.join(pasta, f"log_rede_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
    
    # Log das informações da conexão
    log(f"=== INFORMAÇÕES DA CONEXÃO ===", logfile)
    log(f"Provedor: {info_conexao['provedor']}", logfile)
    log(f"IP Público: {info_conexao['ip']}", logfile)
    log(f"Hostname: {info_conexao['hostname']}", logfile)
    log(f"Tipo IP: {info_conexao['tipo']}", logfile)
    log(f"Confiança: {info_conexao['confianca']}%", logfile)
    log(f"Data/Hora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", logfile)
    log(f"=== INÍCIO DOS TESTES ===\n", logfile)
    
    # Progress bar para todo o processo
    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TaskProgressColumn(),
        console=console
    ) as progress:
        
        # Task principal
        main_task = progress.add_task("🚀 Executando diagnóstico completo...", total=100)
        
        # 1. Resolução DNS (20%)
        progress.update(main_task, description="🔍 Resolvendo DNS...")
        resultados_dns = {}
        for nome, host in DESTINOS.items():
            nome_simples = nome.split(' ', 1)[-1]  # Remove emoji
            ip = resolve_dns(host)
            resultados_dns[f"DNS {nome_simples}"] = ip
            log(f"DNS {nome_simples}: {ip}", logfile)
        
        # Exibir tabela DNS
        console.print("\n")
        console.print(Panel(criar_tabela_dns(resultados_dns), title="🔍 Resolução DNS", border_style="cyan"))
        resultados.update(resultados_dns)
        progress.update(main_task, advance=20)
        
        # 2. Teste de Ping (20%)
        progress.update(main_task, description="📡 Executando testes de ping...")
        console.print("\n")
        console.rule("[bold yellow]📡 Teste de Ping")
        for nome, host in DESTINOS.items():
            nome_simples = nome.split(' ', 1)[-1]  # Remove emoji
            resultado = ping(host)
            
            # Extrair resumo do ping
            ping_info = formatar_ping_resumo(resultado)
            
            # Análise rápida do resultado
            if "100% packet loss" in resultado:
                status_emoji = "❌"
                status_cor = "red"
            elif "0% packet loss" in resultado:
                status_emoji = "✅"
                status_cor = "green"
            else:
                status_emoji = "⚠️"
                status_cor = "yellow"
                
            console.print(f"{status_emoji} [bold]{nome}[/bold] → {ping_info['ip']}")
            
            if ping_info['latencia_avg'] > 0:
                console.print(f"   📊 [cyan]Latência:[/cyan] {ping_info['latencia_avg']:.1f}ms | "
                            f"[cyan]Perda:[/cyan] {ping_info['perda']} | "
                            f"[cyan]Pacotes:[/cyan] {ping_info['pacotes_recebidos']}/{ping_info['pacotes_enviados']}")
            else:
                console.print(f"   📊 [{status_cor}]Sem resposta - {ping_info['perda']} de perda[/{status_cor}]")
            
            console.print()
            resultados[f"Ping {nome_simples}"] = resultado
            log(f"Ping {nome_simples}: {resultado}", logfile)
        progress.update(main_task, advance=20)

        # 3. Traceroute (25%)
        progress.update(main_task, description="🛤️ Executando traceroute...")
        console.print("\n")
        console.rule("[bold blue]🛤️ Traceroute")
        for nome, host in DESTINOS.items():
            nome_simples = nome.split(' ', 1)[-1]  # Remove emoji
            resultado = traceroute(host)
            
            # Extrair resumo do traceroute
            trace_info = formatar_traceroute_resumo(resultado)
            
            console.print(f"🔍 [bold]{nome}[/bold] → {host}")
            console.print(f"   🎯 [cyan]Destino:[/cyan] {trace_info['destino_info'].split(' ')[2] if trace_info['destino_info'] else 'N/A'}")
            
            # Mostrar rota simplificada (primeiros 3, últimos 3 saltos)
            rota = trace_info['rota']
            if len(rota) > 6:
                primeiros = rota[:3]
                ultimos = rota[-3:]
                
                console.print("   🛤️  [cyan]Rota resumida:[/cyan]")
                for hop in primeiros:
                    latencia_str = f"{hop['latencia']}ms" if hop['latencia'] > 0 else "timeout"
                    console.print(f"      {hop['hop']:2d}. {hop['host'][:40]:<40} ({latencia_str})")
                
                if len(rota) > 6:
                    console.print(f"      ... [dim]({len(rota)-6} saltos intermediários)[/dim]")
                
                for hop in ultimos:
                    latencia_str = f"{hop['latencia']}ms" if hop['latencia'] > 0 else "timeout"
                    console.print(f"      {hop['hop']:2d}. {hop['host'][:40]:<40} ({latencia_str})")
            else:
                console.print("   🛤️  [cyan]Rota completa:[/cyan]")
                for hop in rota[:8]:  # Limitar a 8 saltos para não poluir
                    latencia_str = f"{hop['latencia']}ms" if hop['latencia'] > 0 else "timeout"
                    console.print(f"      {hop['hop']:2d}. {hop['host'][:40]:<40} ({latencia_str})")
            
            console.print()
            resultados[f"Traceroute {nome_simples}"] = resultado
            log(f"Traceroute {nome_simples}: {resultado}", logfile)
        progress.update(main_task, advance=25)

        # 4. Análise MTR (20%)
        progress.update(main_task, description="📊 Executando análise MTR...")
        console.print("\n")
        console.rule("[bold magenta]📊 Análise MTR (Perda e Jitter)")
        for nome, host in DESTINOS.items():
            nome_simples = nome.split(' ', 1)[-1]  # Remove emoji
            resultado = mtr_analysis(host)
            
            # Extrair resumo do MTR
            mtr_info = formatar_mtr_resumo(resultado)
            
            console.print(f"📈 [bold]{nome}[/bold] → {host}")
            
            if mtr_info['latencia_final'] > 0:
                console.print(f"   🎯 [cyan]Latência final:[/cyan] {mtr_info['latencia_final']:.1f}ms")
            
            if mtr_info['hops_problematicos']:
                console.print("   ⚠️  [yellow]Saltos com perda detectada:[/yellow]")
                for hop in mtr_info['hops_problematicos'][:5]:  # Mostrar até 5 problemáticos
                    console.print(f"      {hop['hop']:2d}. {hop['host'][:35]:<35} ({hop['perda']:.1f}% perda)")
            else:
                console.print("   ✅ [green]Rota estável - sem perdas significativas[/green]")
            
            console.print()
            resultados[f"MTR {nome_simples}"] = resultado
            log(f"MTR {nome_simples}: {resultado}", logfile)
        progress.update(main_task, advance=20)

        # 5. Speedtest (15%)
        progress.update(main_task, description="🚀 Executando speedtest...")
        console.print("\n")
        console.rule("[bold green]🚀 Speedtest")
        
        resultado = test_speed(server_id=args.speedtest_server)
        if isinstance(resultado, tuple) and len(resultado) == 2 and isinstance(resultado[1], str):
            down, up, ping_ms = resultado[0]
            erro = resultado[1]
            console.print(f"[red]❌ {erro}[/red]")
            speed_result = (0, 0, 0)
        else:
            down, up, ping_ms = resultado
            speed_result = (down, up, ping_ms)
            erro = None
        
        # Exibir tabela de speedtest
        if not erro:
            console.print("\n")
            console.print(Panel(criar_tabela_speedtest(speed_result), title="🚀 Resultado do Speedtest", border_style="green"))
        
        resultados["Speedtest"] = speed_result
        log(f"Speedtest: Download {down} Mbps, Upload {up} Mbps, Latência {ping_ms} ms", logfile)
        progress.update(main_task, advance=15)
        
        # Finalização
        progress.update(main_task, description="✅ Gerando relatórios...", completed=100)

    # Geração de relatórios
    console.print("\n")
    console.rule("[bold cyan]📄 Gerando Relatórios")
    
    with Status("[cyan]Exportando resultados...", spinner="dots"):
        exportar_json(resultados, pasta)
        exportar_csv(resultados, pasta)
        
        resultados_por_destino = extrair_resultados_por_destino(logfile)
        gerar_pdf_resultados_por_destino(resultados_por_destino, resultados["Speedtest"], os.path.join(pasta, "Relatorio_Diagnostico_Rede.pdf"))

        resultados_metricas = extrair_metricas_do_log(logfile)
        gerar_pdf_metricas_por_destino(resultados_metricas, resultados["Speedtest"], pasta, "Relatorio_Diagnostico_Rede_Metricas.pdf")
    
    # Mensagem final
    final_panel = Panel(
        Align.center("✅ [bold green]Diagnóstico Concluído![/bold green]\n🎯 Todos os relatórios foram gerados com sucesso!"), 
        border_style="green",
        padding=(1, 2)
    )
    console.print(final_panel)

if __name__ == "__main__":
    main()
